import re
import sys
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(r"C:\Users\XiaoyanXiong\Desktop\NIDSTransformer")
sys.path.insert(0, str(ROOT / "defense_build"))
from qa_content import CATEGORIES

path = ROOT / "Defense_Professor_Questions_and_Answers_Xiaoyan_Xiong.docx"
doc = Document(path)
questions = [
    p for p in doc.paragraphs
    if p.style.name == "Heading 2" and re.match(r"Q[0-9]{3}", p.text)
]
numbers = [int(re.match(r"Q([0-9]{3})", p.text).group(1)) for p in questions]
text = "\n".join(p.text for p in doc.paragraphs)
expected = sum(len(category[2]) for category in CATEGORIES)
with zipfile.ZipFile(path) as package:
    zip_error = package.testzip()

print("categories", len(CATEGORIES))
print("expected_questions", expected)
print("question_headings", len(questions))
print("sequential_numbering", numbers == list(range(1, expected + 1)))
print("suggested_answer_labels", text.count("Suggested answer (simple English):"))
print("chinese_logic_labels", text.count("中文回答逻辑："))
print("boundary_labels", text.count("Evidence boundary / 避免过度表述："))
print("english_tokens", len(re.findall(r"[A-Za-z][A-Za-z0-9_-]*", text)))
print("chinese_chars", len(re.findall(r"[\u4e00-\u9fff]", text)))
print("codex_tokens", text.count(":codex-"))
print("placeholder_tokens", len(re.findall(r"TODO|TBD|PLACEHOLDER", text, re.I)))
print("zip_test", zip_error or "OK")
print("file_bytes", path.stat().st_size)
