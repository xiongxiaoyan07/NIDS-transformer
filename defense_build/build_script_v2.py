import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(r'C:\Users\XiaoyanXiong\Desktop\NIDSTransformer')
OUT=ROOT/'Defense_Speech_Script_Xiaoyan_Xiong_Revised.docx'
slides=json.loads((ROOT/'defense_build'/'notes_v2.json').read_text(encoding='utf-8'))

doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(.72); sec.bottom_margin=Inches(.72)
sec.left_margin=Inches(.82); sec.right_margin=Inches(.82)

normal=doc.styles['Normal']
normal.font.name='Arial'; normal.font.size=Pt(11)
normal.paragraph_format.space_after=Pt(7); normal.paragraph_format.line_spacing=1.15
for sty,size,color,before,after in [
    ('Heading 1',16,'1677B8',14,7),('Heading 2',13,'1677B8',10,5)]:
    s=doc.styles[sty]; s.font.name='Arial'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('MSc Thesis Defense - Revised Speaker Script'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(25); r.font.color.rgb=RGBColor.from_string('111318')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('A Hierarchical Time-Aware Transformer for Flow-Level Network Intrusion Detection'); r.font.name='Arial'; r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string('1677B8')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Xiaoyan Xiong | Planned speaking time: about 25 minutes'); r.italic=True; r.font.name='Arial'; r.font.size=Pt(11)

doc.add_heading('How to use this script',level=1)
for text in [
    'The wording uses short sentences and common English words. Do not try to speak faster than your natural pace.',
    'The last sentence of each slide is a transition to the next slide. Keep it, because it makes the presentation sound connected.',
    'When a chart appears, point to the relevant bar or number before saying the result.',
    'The planned slide times add up to about 27 minutes on paper; natural compression during rehearsal should bring the talk close to 24-25 minutes.',
    'Use the term alert-associated class when you want to avoid implying independently verified attack ground truth.'
]:
    doc.add_paragraph(text,style='List Bullet')

doc.add_heading('Timing overview',level=1)
table=doc.add_table(rows=1,cols=3)
table.alignment=WD_TABLE_ALIGNMENT.CENTER
table.style='Table Grid'
hdr=table.rows[0].cells
hdr[0].text='Slide'; hdr[1].text='Topic'; hdr[2].text='Target time'
for i,item in enumerate(slides,1):
    cells=table.add_row().cells
    cells[0].text=str(i); cells[1].text=item['title'].split('. ',1)[-1]; cells[2].text=item['time']
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_after=Pt(2)
            for run in p.runs: run.font.name='Arial'; run.font.size=Pt(9.5)
for cell in table.rows[0].cells:
    shading=OxmlElement('w:shd'); shading.set(qn('w:fill'),'EAF6FC'); cell._tc.get_or_add_tcPr().append(shading)
    for run in cell.paragraphs[0].runs: run.bold=True

doc.add_page_break()
doc.add_heading('Detailed slide-by-slide script',level=1)
for i,item in enumerate(slides,1):
    doc.add_heading(f"Slide {i}: {item['title'].split('. ',1)[-1]} ({item['time']})",level=2)
    p=doc.add_paragraph(item['script'])
    p.paragraph_format.keep_together=False
    cue=doc.add_paragraph()
    cue.paragraph_format.space_after=Pt(8)
    r=cue.add_run('Practice cue: '); r.bold=True; r.font.color.rgb=RGBColor.from_string('1677B8')
    if i in (6,8,11,12,13,14): cue.add_run('Point to the figure or highlighted metric; pause for one second before moving on.')
    elif i==15: cue.add_run('Read only the range and pooled result aloud; do not read every table cell.')
    elif i==17: cue.add_run('Look at the committee, slow down, and finish clearly before inviting questions.')
    else: cue.add_run('Keep eye contact and use the final sentence as the transition to the next slide.')

doc.add_heading('Likely defense questions and simple answers',level=1)
qas=[
('Why did you use Suricata-derived labels?', 'They gave me a deterministic way to label a large enterprise capture. However, I treat them as weak labels. My claim is about alert-associated behaviour, not perfect attack ground truth.'),
('Why did you separate Stage 1 and Stage 2?', 'The separation follows the traffic hierarchy and makes the experiments easier to interpret. Stage 1 learns the flow representation. Stage 2 tests the additional value of historical context. End-to-end training is possible future work.'),
('Is Stage 2 really causal?', 'It is causal in stable flow-start order. It excludes future embeddings and all historical labels. It is not yet strict streaming causality because an earlier long flow may not have completed before the target starts.'),
('Why was source-host context the best?', 'It was the best observed relation in this capture. Repeated activity from one source may preserve scanning or retry behaviour. Because the labels are binary, I present this as a possible explanation, not a proven mechanism.'),
('Why use PR-AUC?', 'The positive class is rare. PR-AUC focuses on ranking positive examples and is more informative than accuracy alone under strong imbalance.'),
('Why not report statistical significance?', 'The experiments use one fixed seed because of the available compute budget. I therefore treat small differences as descriptive and do not claim statistical significance.'),
('Can the model generalise to other networks?', 'The external tests show promising transfer to CICIDS2017 windows and another company server. However, the performance varies, so the evidence is preliminary and not universal.'),
('Why use a Transformer instead of only an LSTM?', 'Self-attention allows distant packets or flows to interact directly and supports parallel training. The controlled results also show a better error balance for the proposed model in this study.'),
]
for q,a in qas:
    p=doc.add_paragraph(); r=p.add_run(q); r.bold=True; r.font.color.rgb=RGBColor.from_string('1677B8')
    doc.add_paragraph(a)

doc.save(OUT)
print(OUT)
