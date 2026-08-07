from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from qa_content import CATEGORIES


ROOT = Path(r"C:\Users\XiaoyanXiong\Desktop\NIDSTransformer")
OUT = ROOT / "Defense_Professor_Questions_and_Answers_Xiaoyan_Xiong.docx"

BLUE = "1677B8"
DARK = "102A43"
MUTED = "52606D"
RED = "9B1C1C"
GOLD = "7A5A00"
LIGHT_BLUE = "E8EEF5"
BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = BODY_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def add_label_paragraph(doc, label, text, label_color=BLUE, text_color=None, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    r = p.add_run(label)
    set_run_font(r, size=11, bold=True, color=label_color)
    r = p.add_run(text)
    set_run_font(r, size=11, color=text_color)
    return p


def add_real_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, bold=True, color=BLUE)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=11)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# compact_reference_guide preset with one named override:
# Microsoft YaHei is used for East Asian glyphs.
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
):
    style = doc.styles[name]
    style.font.name = BODY_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_style_name in ("List Bullet", "List Number"):
    style = doc.styles[list_style_name]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.paragraph_format.space_after = Pt(0)
r = header.add_run("MSc Thesis Defense | Professor Q&A Bank")
set_run_font(r, size=8.5, color=MUTED)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.paragraph_format.space_before = Pt(0)
r = footer.add_run("Page ")
set_run_font(r, size=8.5, color=MUTED)
add_field(footer, "PAGE")

# Editorial-cover pattern.
for _ in range(3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run("MSc Thesis Defense Preparation")
set_run_font(r, size=12, bold=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Professor Questions and Answers")
set_run_font(r, size=28, bold=True, color=DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("教授可能提问与答辩参考回答")
set_run_font(r, size=20, bold=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
r = p.add_run("A Hierarchical Time-Aware Transformer for Flow-Level Network Intrusion Detection")
set_run_font(r, size=13.5, color=DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
r = p.add_run("Xiaoyan Xiong")
set_run_font(r, size=13, bold=True, color=DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("101 questions | Simple spoken English answers | Chinese reasoning | Evidence boundaries")
set_run_font(r, size=10.5, italic=True, color=MUTED)

# Flatten and number all questions.
numbered = []
question_number = 0
for category_en, category_zh, items in CATEGORIES:
    category_items = []
    for item in items:
        question_number += 1
        record = {
            "number": question_number,
            "priority": item[0],
            "q_en": item[1],
            "q_zh": item[2],
            "answer_en": item[3],
            "answer_zh": item[4],
            "boundary": item[5],
        }
        numbered.append(record)
        category_items.append(record)
    # Store numbered category data in-place for the render loop.
    items[:] = category_items

doc.add_page_break()
doc.add_heading("How to use this question bank / 使用方法", level=1)
for text in (
    "先练习前面的高频问题，再按自己最不熟悉的主题进行第二轮练习。不要尝试逐字背诵全部 101 个答案。",
    "英文回答采用较短句子和常用词。答题时先给直接结论，再给一到两个论文证据，最后说明证据边界。",
    "标记 A 的问题属于高频或高风险问题；标记 B 的问题属于深入追问。",
    "如果教授的问题超出论文实验范围，应明确说没有直接测试，并说明你会如何设计后续实验。",
    "所有 class 1 标签应表述为 alert-associated class，而不是 independently verified attack ground truth。",
):
    add_real_bullet(doc, text)

doc.add_heading("A safe four-step answer structure / 安全答题结构", level=1)
for text in (
    "1. Direct answer: Yes, no, or the main conclusion.",
    "2. Evidence: Give one method detail, comparison, or exact number from the thesis.",
    "3. Boundary: State what the experiment did not establish.",
    "4. Stop: Do not add a stronger claim after a correct answer.",
):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    # Avoid a duplicated visual numeral by stripping the textual prefix.
    content = text.split(". ", 1)[1]
    r = p.add_run(content)
    set_run_font(r, size=11)

doc.add_heading("Useful rescue phrases / 不确定时的安全英文表达", level=1)
for text in (
    "Thank you. I would separate this into two points.",
    "I did not test that directly, so I do not want to overclaim.",
    "In the evaluated setting, the result supports this explanation, but it does not prove the mechanism.",
    "You are right. This is a limitation of the current study.",
    "The exact comparison is valid only because these models share the same test manifest.",
    "I would test that with repeated seeds and a parameter-matched control.",
    "The model predicts an alert-associated flow, not independently verified malicious traffic.",
):
    add_real_bullet(doc, text)

doc.add_heading("Numbers you should memorise / 建议记住的核心数字", level=1)
for text in (
    "Source data: 8.62 GB PCAP; 11,750,727 packet rows; 421,898 flow rows.",
    "Labels: 16,644 positive flows; 3.945% positive; imbalance ratio 24.35:1.",
    "Split: 295,329 train; 42,189 validation; 84,380 test flows.",
    "Stage 1 position+time: macro-F1 0.8882; PR-AUC 0.8564.",
    "Stage 1 complete Scheme C: macro-F1 0.8954; PR-AUC 0.8763; FP 267; FN 293.",
    "Stage 2 source-host: macro-F1 0.9275; PR-AUC 0.9345; FPR 0.5429%.",
    "CICIDS2017 pooled: 100,000 flows; 15.353% positive; macro-F1 0.8966.",
    "Company cross-server: 38,513 flows; 3.285% positive; macro-F1 0.9026; PR-AUC 0.8771.",
):
    add_real_bullet(doc, text)

doc.add_page_break()
doc.add_heading("Twenty questions to practise first / 优先练习的 20 个问题", level=1)
top_numbers = [1, 6, 8, 9, 14, 15, 22, 25, 26, 28, 31, 33, 34, 38, 41, 44, 47, 55, 69, 96]
by_number = {item["number"]: item for item in numbered}
for number in top_numbers:
    item = by_number[number]
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(f"Q{number:03d}: {item['q_en']} / {item['q_zh']}")
    set_run_font(r, size=11)

doc.add_heading("Question-bank categories / 问题分类", level=1)
for category_en, category_zh, items in CATEGORIES:
    add_real_bullet(doc, f"{category_en} / {category_zh} ({len(items)} questions)")

# Detailed questions.
for category_en, category_zh, items in CATEGORIES:
    doc.add_page_break()
    doc.add_heading(category_en, level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(category_zh)
    set_run_font(r, size=12, bold=True, color=MUTED)
    for item in items:
        level = "High priority / 高频高风险" if item["priority"] == "A" else "Deep follow-up / 深入追问"
        h = doc.add_heading(f"Q{item['number']:03d} [{item['priority']}] {item['q_en']}", level=2)
        h.paragraph_format.keep_with_next = True
        add_label_paragraph(doc, "中文问题：", item["q_zh"], label_color=BLUE, keep=True)
        add_label_paragraph(doc, "Suggested answer (simple English): ", item["answer_en"], label_color=DARK)
        add_label_paragraph(doc, "中文回答逻辑：", item["answer_zh"], label_color=BLUE)
        add_label_paragraph(doc, "Evidence boundary / 避免过度表述：", item["boundary"], label_color=RED, text_color=RED)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(f"Priority: {level}")
        set_run_font(r, size=9.5, italic=True, color=GOLD if item["priority"] == "A" else MUTED)

doc.add_page_break()
doc.add_heading("Final rehearsal checklist / 最终练习清单", level=1)
for text in (
    "能够准确说出两个 research questions 的论文原文。",
    "能够在 30 秒内解释 packet、flow、Stage 1 和 Stage 2。",
    "能够解释 weak labels，并明确 label 0 不等于 verified benign。",
    "能够区分 start-ordered causality 与 completion-time streaming causality。",
    "能够解释为什么不同 test manifest 的结果不能直接比较。",
    "能够解释 macro-F1、PR-AUC、class-1 F1、FPR、FP 和 FN 的角色。",
    "能够说出 Stage 1、Stage 2 和两个 external evaluations 的核心数字。",
    "能够主动承认 one seed、no capacity-matched control 和 incomplete end-to-end latency。",
    "遇到未测试的问题时，先说没有直接证据，再提出合理实验，而不是猜测。",
):
    add_real_bullet(doc, text)

doc.core_properties.title = "MSc Thesis Defense Professor Questions and Answers"
doc.core_properties.subject = "A Hierarchical Time-Aware Transformer for Flow-Level Network Intrusion Detection"
doc.core_properties.author = "Xiaoyan Xiong"
doc.core_properties.keywords = "thesis defense, NIDS, Transformer, questions, answers"
doc.save(OUT)
print(OUT)
