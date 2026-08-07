import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\XiaoyanXiong\Desktop\NIDSTransformer")
OUT = ROOT / "Defense_Speech_Script_Xiaoyan_Xiong_Chinese.docx"
SLIDES = json.loads((ROOT / "defense_build" / "notes_zh.json").read_text(encoding="utf-8"))

BLUE = "1677B8"
DARK = "102A43"
MUTED = "52606D"
LIGHT_BLUE = "E8EEF5"
WHITE = "FFFFFF"
BODY_FONT = "Arial"
EAST_ASIA_FONT = "Microsoft YaHei"


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), BODY_FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), BODY_FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)


def add_real_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# compact_reference_guide preset; named Chinese-font override for East Asian glyphs.
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
):
    style = doc.styles[name]
    style.font.name = BODY_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

# Real list style with preset spacing and indents.
bullet_style = doc.styles["List Bullet"]
bullet_style.font.name = BODY_FONT
bullet_style.font.size = Pt(11)
bullet_style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
bullet_style.paragraph_format.left_indent = Inches(0.375)
bullet_style.paragraph_format.first_line_indent = Inches(-0.188)
bullet_style.paragraph_format.space_after = Pt(4)
bullet_style.paragraph_format.line_spacing = 1.25

# Quiet running header and footer.
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_p.paragraph_format.space_after = Pt(0)
r = header_p.add_run("硕士论文答辩 | 中文演讲稿")
set_run_font(r, size=8.5, color=MUTED)

footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(0)
r = footer_p.add_run("Xiaoyan Xiong")
set_run_font(r, size=8.5, color=MUTED)

# editorial_cover first-page pattern, adapted for a practical defense guide.
for _ in range(3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run("硕士论文答辩")
set_run_font(r, size=12, bold=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("中文演讲稿")
set_run_font(r, size=28, bold=True, color=DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("A Hierarchical Time-Aware Transformer for")
set_run_font(r, size=14, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
r = p.add_run("Flow-Level Network Intrusion Detection")
set_run_font(r, size=14, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
r = p.add_run("Xiaoyan Xiong")
set_run_font(r, size=13, bold=True, color=DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("目标演讲时长：约 25 分钟 | 对应 17 页修改版 PPT")
set_run_font(r, size=10.5, italic=True, color=MUTED)

doc.add_page_break()
doc.add_heading("使用说明", level=1)
for text in (
    "本稿逐页对应修改版 PPT。英文专业术语被保留，便于演讲时与页面内容保持一致。",
    "每页最后一句是通向下一页的过渡语，建议保留，不要在换页时停顿太久。",
    "图表页不需要读出全部数字，只强调页面上突出显示的指标和对结论最重要的比较。",
    "建议按自己的自然语速练习两次，并在第 8、9、14 和 15 页预留指图和停顿时间。",
    "涉及标签时，建议使用“alert-associated class”或“告警关联类别”，避免把标签表述为独立验证的攻击真值。",
):
    add_real_bullet(doc, text)

doc.add_heading("时间总览", level=1)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = ("页码", "主题", "建议时间")
for idx, label in enumerate(headers):
    cell = table.rows[0].cells[idx]
    cell.text = label
    set_cell_fill(cell, LIGHT_BLUE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for run in cell.paragraphs[0].runs:
        set_run_font(run, size=9.5, bold=True, color=DARK)
for idx, item in enumerate(SLIDES, start=1):
    cells = table.add_row().cells
    values = (str(idx), item["title_zh"], item["time"])
    for cidx, value in enumerate(values):
        cells[cidx].text = value
        cells[cidx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[cidx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if cidx != 1 else WD_ALIGN_PARAGRAPH.LEFT
        cells[cidx].paragraphs[0].paragraph_format.space_after = Pt(0)
        for run in cells[cidx].paragraphs[0].runs:
            set_run_font(run, size=9.2)
set_repeat_table_header(table.rows[0])
set_table_geometry(table, [900, 6960, 1500])

doc.add_page_break()
doc.add_heading("逐页中文演讲稿", level=1)
for idx, item in enumerate(SLIDES, start=1):
    heading = doc.add_heading(
        f"第 {idx} 页：{item['title_zh']} / {item['title_en']}（{item['time']}）",
        level=2,
    )
    heading.paragraph_format.keep_with_next = True
    p = doc.add_paragraph(item["script"])
    p.paragraph_format.keep_together = False
    p.paragraph_format.widow_control = True
    cue = doc.add_paragraph()
    cue.paragraph_format.space_before = Pt(2)
    cue.paragraph_format.space_after = Pt(8)
    cue.paragraph_format.keep_together = True
    r = cue.add_run("练习提示：")
    set_run_font(r, size=10, bold=True, color=BLUE)
    if idx in (6, 8, 11, 12, 13, 14):
        cue_text = "先指向图中的流程或重点指标，再说结论；切换下一页前停顿一秒。"
    elif idx == 15:
        cue_text = "只读 positive proportion 的范围和 pooled result，不需要逐个读完整表格。"
    elif idx == 17:
        cue_text = "放慢语速，看向答辩委员会，清楚说完最后一句后再邀请提问。"
    else:
        cue_text = "保持目光交流，并把最后一句作为切换到下一页的过渡。"
    r = cue.add_run(cue_text)
    set_run_font(r, size=10, italic=True, color=MUTED)

doc.add_heading("常见答辩问题与中文参考答案", level=1)
qas = [
    (
        "为什么使用 Suricata 生成的标签？",
        "Suricata 为大型企业网络数据提供了一种可重复、可执行的标注方式。但是，我把它看作 weak supervision。论文的结论针对 alert-associated behaviour，而不是把这些标签当作完美的 attack ground truth。",
    ),
    (
        "为什么把模型分成 Stage 1 和 Stage 2？",
        "这种划分与网络流量的层次结构一致，也让实验更容易解释。Stage 1 学习 flow representation，Stage 2 单独检验 historical context 的额外价值。未来可以进一步研究 end-to-end training。",
    ),
    (
        "Stage 2 是否真正满足 causal？",
        "它在稳定的 flow-start order 中是 causal 的，不使用 future embedding，也不使用历史 label。但它还不是严格的 streaming causality，因为一个更早开始的长 flow 在 target flow 开始时可能还没有结束。",
    ),
    (
        "为什么 source-host context 的效果最好？",
        "这是当前 capture 中观察到的最佳 relation。同一 source 的重复活动可能保留扫描、重试或相关连接模式。但由于标签是 binary 的，我把它解释为可能原因，而不是已经证明的机制。",
    ),
    (
        "为什么使用 PR-AUC？",
        "因为 positive class 很少。PR-AUC 更关注模型对 positive example 的排序能力，在强类别不平衡时通常比单独报告 accuracy 更有信息。",
    ),
    (
        "为什么没有报告统计显著性？",
        "受计算资源限制，实验使用一个固定 seed。因此，我把小的性能差异看作描述性结果，不把它解释为 statistical significance。",
    ),
    (
        "模型是否能够泛化到其他网络？",
        "CICIDS2017 和另一台 company server 的结果说明模型具有一定的 transfer ability。但是，不同窗口的结果存在变化，而且外部来源数量有限，所以这个结论是初步的，不是 universal generalisation。",
    ),
    (
        "为什么使用 Transformer，而不是只使用 LSTM？",
        "Self-attention 可以让距离较远的 packet 或 flow 直接交互，也适合并行训练。在本研究的受控比较中，提出的方法也表现出更好的综合 error balance。",
    ),
]
for question, answer in qas:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(question)
    set_run_font(r, size=11, bold=True, color=BLUE)
    p = doc.add_paragraph(answer)
    p.paragraph_format.keep_together = True

doc.core_properties.title = "硕士论文答辩中文演讲稿"
doc.core_properties.subject = "A Hierarchical Time-Aware Transformer for Flow-Level Network Intrusion Detection"
doc.core_properties.author = "Xiaoyan Xiong"
doc.core_properties.keywords = "NIDS, Transformer, packet, flow, defense script"
doc.save(OUT)
print(OUT)
